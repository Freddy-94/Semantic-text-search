"""
Author: Alfredo Bernal Luna
Date: 18/05/24
Name: extract_text.py
Principal functions:
    1. extract_text_from_pdf
    2. extract_text_from_docx
"""

import fitz  
from docx import Document

def extract_text_from_pdf(pdf_path, main_title="", top_margin=50, bottom_margin=50):
    """
        Function to extract text from a PDF file. Notice that this function helps us fo extract text for all 3 pdf files
    """
    try:
        document = fitz.open(pdf_path)
        all_text = []
        header_texts = set()
        footer_texts = set()
        for page_num in range(len(document)):
            page = document.load_page(page_num)
            page_height = page.rect.height
            # Get text blocks
            blocks = page.get_text("dict")["blocks"]
            page_text = ""
            for block in blocks:
                bbox = block["bbox"]
                block_top = bbox[1]
                block_bottom = bbox[3]
                block_text = ""
                if "lines" in block:
                    for line in block["lines"]:
                        if "spans" in line:
                            line_text = ""
                            for span in line["spans"]:
                                line_text += span["text"]
                            block_text += line_text.strip() #+ " "
                # Exclude text blocks that match the main title as per it's repeating with multiple letters
                if block_text == main_title:
                    continue
                # Exclude text blocks in the top and bottom margins
                if block_top > top_margin and block_bottom < (page_height - bottom_margin):
                    page_text += block_text + " " # Add a space to separate blocks
                else:
                    # Add to header or footer sets
                    if block_top <= top_margin:
                        header_texts.add(block_text)
                    elif block_bottom >= (page_height - bottom_margin):
                        footer_texts.add(block_text)
            all_text.append(page_text.strip())
        # Remove headers and footers from the collected text
        final_text = []
        for page_text in all_text:
            for header in header_texts:
                if page_text.startswith(header):
                    page_text = page_text[len(header):].strip()
            for footer in footer_texts:
                if page_text.endswith(footer):
                    page_text = page_text[:-len(footer)].strip()
            final_text.append(page_text)                    
        return "\n".join(final_text).strip()
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def extract_text_from_docx(docx_path):
    """
    Function to extract text from a DOCX file, ignoring headers and footers, and excluding typical email headers.
    """
    try:
        document = Document(docx_path)
        all_text = []
        header_texts = set()
        footer_texts = set()
        
        # Extract headers and footers from each section
        for section in document.sections:
            header = section.header
            footer = section.footer
            
            for paragraph in header.paragraphs:
                header_texts.add(paragraph.text.strip())
            
            for paragraph in footer.paragraphs:
                footer_texts.add(paragraph.text.strip())
        
        in_email_body = False
        email_headers = ["message-id:", "date:", "from:", "to:", "subject:", "cc:", "mime-version:",\
                         "content-type:", "content-transfer-encoding:", "bcc:", "x-from", "x-to", "x-cc:",\
                         "x-bcc", "x-folder", "x-origin", "x-filename"]

        for paragraph in document.paragraphs:
            paragraph_text = paragraph.text.strip()
            
            # Identify the start of the email body by looking for typical email headers
            if any(paragraph_text.lower().startswith(header) for header in email_headers):
                continue  # Skip typical email header lines

            # Skip headers and footers
            if paragraph_text in header_texts or paragraph_text in footer_texts:
                continue

            # Include paragraph text
            all_text.append(paragraph_text)
        
        return "\n".join(all_text).strip()
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def main():
    #pdf_path = 'Salim v Mitchell - Deposition of Jose Rodriguez.pdf'  
    #text = extract_text_from_pdf(pdf_path, main_title='',top_margin=50, bottom_margin=50)
    docx_path = 'Message-ID_ _29870958.docx'
    text = extract_text_from_docx(docx_path)
    if text:
        print(text)
    else:
        print("No text was extracted.")

if __name__ == '__main__':
    main()
